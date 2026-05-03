# -*- coding: utf-8 -*-
"""Пересборка п. 3.2: ГОСТ (Times New Roman 14, 1,5, по ширине, абзац 1,25 см)."""
import pathlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = pathlib.Path(__file__).resolve().parent
DOC = sorted(ROOT.glob("*.docx"))[3]
FIG_DIR = ROOT / "doc_figures"

BUL = "\u25aa"  # маленький квадрат, как в примерах


def gost_base(p, justify=True):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.first_line_indent = Cm(1.25)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_body(doc, text, *, bullet=False, justify=True):
    p = doc.add_paragraph()
    gost_base(p, justify=justify)
    t = f"{BUL} {text}" if bullet else text
    r = p.add_run(t)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(14)
    return p


def add_captionFigure(doc, n: int, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run(f"Рисунок {n} – {title}")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(14)
    r.italic = True
    return p


def add_heading_3_main(doc) -> None:
    """Один раздел 3 (без второго «3. …» внутри) — заголовок главы 3 по ГОСТ."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("3. Проектирование пользовательского интерфейса")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(14)
    r.bold = True


def add_picture_if_exists(doc, n: int, max_w_in=5.5):
    path = FIG_DIR / f"fig{n:02d}.png"
    if not path.is_file():
        # Если файлов нет, не добавляем техническую заглушку в текст документа.
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(max_w_in))


def apply_gost_margins(doc: Document) -> None:
    """Поля: левое 30 мм, правое 15 мм, верх/низ 20 мм (типовой вариант для пояснительных записей)."""
    for sec in doc.sections:
        sec.left_margin = Cm(3)
        sec.right_margin = Cm(1.5)
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)


def main():
    doc = Document(str(DOC))
    apply_gost_margins(doc)
    # Срезать весь раздел 3: от «3. Проектирование пользовательского...» до конца.
    cut_from = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "3. Проектирование пользовательского интерфейса":
            cut_from = i
            break
    if cut_from is None:
        raise SystemExit("Не найдено начало раздела 3 для замены.")

    body = doc._body._element
    for p in list(doc.paragraphs[cut_from:]):
        body.remove(p._p)

    # Только раздел 3: кратко и по ключевым экранам для объема ~2 страницы
    add_heading_3_main(doc)

    INTRO_3 = [
        "Проектирование пользовательского интерфейса выполнялось в рамках единой концепции визуальной и функциональной целостности веб-приложения. При разработке учитывались требования к понятной навигации, читаемости контента, предсказуемости пользовательских сценариев и согласованности экранов для разных ролей системы.",
        "Независимо от типа страницы (публичной, авторизационной или административной) используются единые стилистические принципы: светлая базовая тема, акцентная цветовая палитра, унифицированная типографика и адаптивная верстка для различных разрешений экрана.",
    ]
    for para in INTRO_3:
        add_body(doc, para, bullet=False)

    # Ключевые экраны: главная, регистрация, админка
    blocks = [
        (
            False,
            "Главная страница выступает точкой входа в приложение и направляет пользователя к базовым сценариям: просмотру туров, переходу к подробной информации и дальнейшему бронированию. Компоновка главного экрана выполнена так, чтобы ключевые действия были доступны без избыточной навигации (см. рисунок 1).",
        ),
        (
            False,
            "Экран регистрации обеспечивает ввод и первичную валидацию учетных данных пользователя. Форма организована в понятной последовательности полей, а текстовые подсказки и сообщения об ошибках помогают снизить количество некорректных действий на этапе создания учетной записи (см. рисунок 2).",
        ),
        (
            False,
            "Административная часть вынесена в отдельный интерфейс и предназначена для управления турами, пользователями и служебными процессами платформы. Разделение публичного и административного контуров повышает удобство сопровождения и обеспечивает более прозрачную структуру ролей в системе (см. рисунок 3).",
        ),
    ]

    captions = [
        "Главная страница веб-приложения",
        "Экран регистрации пользователя",
        "Административный интерфейс",
    ]

    for idx, (is_bull, text) in enumerate(blocks, start=1):
        add_body(doc, text, bullet=is_bull)
        add_picture_if_exists(doc, idx)
        add_captionFigure(doc, idx, captions[idx - 1] if idx <= len(captions) else f"Рисунок {idx}")

    doc.save(str(DOC))
    print("OK:", DOC.name)


if __name__ == "__main__":
    main()
