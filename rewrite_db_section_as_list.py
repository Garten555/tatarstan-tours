# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import pathlib


def set_run_font(run, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = bold


def add_paragraph(doc: Document, text: str, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_bullet_line(doc: Document, text: str):
    # Для Word-документа используем "— " как маркер списка
    p = add_paragraph(doc, f"— {text}", indent=True)
    return p


def main():
    root = pathlib.Path(__file__).resolve().parent
    target = root / "ПРАКТИЧЕСКАЯ_ЧАСТЬ_3_ПУНКТА.docx"
    doc = Document(str(target))

    # Ищем границы раздела 2
    start = None
    end = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "2. Проектирование базы данных":
            start = i
            continue
        if start is not None and t.startswith("3. "):
            end = i
            break

    if start is None:
        raise SystemExit("Не найден заголовок раздела 2.")
    if end is None:
        end = len(doc.paragraphs)

    body = doc._body._element
    for p in list(doc.paragraphs[start + 1:end]):
        body.remove(p._p)

    # Вставляем новый текст сразу после заголовка раздела 2
    # Самый простой безопасный способ: добавляем в конец, потом переносим блок.
    new_block = [
        ("p", "При проектировании базы данных выделялись следующие сущности (с указанием назначения каждой таблицы):"),
        ("b", "profiles — хранит учетные и профильные данные пользователя, роль, статус и признаки модерации."),
        ("b", "cities — справочник городов и локаций, используемый в карточках туров и фильтрации."),
        ("b", "tours — каталог туров: название, описание, параметры маршрута, даты, стоимость, вместимость и статус публикации."),
        ("b", "bookings — фиксирует факт бронирования тура пользователем и статусы заявки/оплаты."),
        ("b", "booking_attendees — хранит состав участников по каждой брони и персональные данные путешественников."),
        ("b", "travel_diaries — содержит пользовательские дневники поездок, текстовый контент и геоданные публикаций."),
        ("b", "achievements — хранит достижения и бейджи пользователя, условия и дату получения."),
        ("b", "support_sessions — описывает сессии обращений в поддержку (контекст диалога, состояние, сроки)."),
        ("b", "chat_messages — хранит сообщения внутри сессий поддержки с привязкой к отправителю."),
        ("b", "notifications — содержит события и уведомления для пользователя, включая признак прочтения."),
        ("p", "Для наглядного представления структуры данных и взаимосвязей между сущностями используется ERD-диаграмма."),
        ("p", "[ВСТАВИТЬ ERD-ДИАГРАММУ БАЗЫ ДАННЫХ]"),
        ("p", "На рисунке X представлена ERD-диаграмма базы данных разработанного веб-приложения."),
        ("p", "Диаграмма подтверждает корректность выбранной модели: таблицы пользовательского контура, каталога туров, бронирований, дневников и служебных коммуникаций связаны в единую реляционную структуру, обеспечивающую целостность и расширяемость системы."),
    ]

    inserted = []
    for kind, text in new_block:
        if kind == "b":
            inserted.append(add_bullet_line(doc, text))
        else:
            inserted.append(add_paragraph(doc, text, indent=True))

    # Переносим вставленные абзацы под заголовок раздела 2
    anchor = doc.paragraphs[start]._p
    parent = anchor.getparent()
    anchor_index = parent.index(anchor)
    for p in inserted:
        el = p._p
        parent.remove(el)
        parent.insert(anchor_index + 1, el)
        anchor_index += 1

    doc.save(str(target))
    print("updated:", target)


if __name__ == "__main__":
    main()

